"""Generate the EventCart QA application flow Word document from Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
CODE_FILL = "F4F6F9"
BORDER = "B7C9DD"


def set_cell_shading(cell, fill: str) -> None:
    """Apply a background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    """Apply Word table cell margins in DXA."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    """Apply fixed table geometry with explicit grid and cell widths."""
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    """Apply compact reference guide styling tokens."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "EventCart QA Application Flow Guide"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Living QA and onboarding document"
    footer.style = styles["Normal"]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    """Add the first-page title block."""
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("EventCart QA And New Joiner Application Flow Guide")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "A living handoff guide for QA execution, local verification, and new joiner onboarding."
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(meta, [1900, CONTENT_WIDTH_DXA - 1900])
    rows = (
        ("Project", "EventCart real-time e-commerce order platform"),
        ("Current scope", "Catalog, cart, order, inventory, Redis idempotency, Kafka order status flow"),
        ("Last updated", "2026-08-02"),
    )
    for row_index, (label, value) in enumerate(rows):
        meta.cell(row_index, 0).text = label
        meta.cell(row_index, 1).text = value
        for cell in meta.rows[row_index].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
        set_cell_shading(meta.cell(row_index, 0), HEADER_FILL)
        meta.cell(row_index, 0).paragraphs[0].runs[0].bold = True


def parse_inline_code(paragraph, text: str) -> None:
    """Add text with simple inline-code run handling."""
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(9.5)


def add_code_block(doc: Document, lines: list[str]) -> None:
    """Add a code block using monospaced paragraphs and light shading."""
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Inches(0.15)
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(8.5)
        p_pr = paragraph._p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), CODE_FILL)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a Markdown table with explicit fixed-width geometry."""
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and all(set(cell) <= {"-", " "} for cell in rows[1]) else rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for index, cell_text in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = cell_text
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row in body:
        cells = table.add_row().cells
        for index, cell_text in enumerate(row):
            cells[index].text = cell_text
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    col_count = len(header)
    if col_count == 2:
        widths = [2200, CONTENT_WIDTH_DXA - 2200]
    elif col_count == 3:
        widths = [1900, 3100, CONTENT_WIDTH_DXA - 1900 - 3100]
    elif col_count == 4:
        widths = [1800, 2100, 2300, CONTENT_WIDTH_DXA - 6200]
    else:
        base = CONTENT_WIDTH_DXA // col_count
        widths = [base] * col_count
        widths[-1] = CONTENT_WIDTH_DXA - sum(widths[:-1])
    set_table_width(table, widths)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def markdown_table_row(line: str) -> list[str]:
    """Parse one simple Markdown pipe-table row."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_to_docx(doc: Document, markdown: str) -> None:
    """Convert a practical subset of Markdown into the Word document."""
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("# "):
            index += 1
            continue

        if line.startswith("|"):
            table_rows.append(markdown_table_row(line))
            index += 1
            if index == len(lines) or not lines[index].startswith("|"):
                add_markdown_table(doc, table_rows)
                table_rows = []
            continue

        if not line.strip():
            index += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            index += 1
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            parse_inline_code(paragraph, line[2:].strip())
            index += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered_match:
            paragraph = doc.add_paragraph(style="List Number")
            parse_inline_code(paragraph, numbered_match.group(1).strip())
            index += 1
            continue

        paragraph = doc.add_paragraph()
        parse_inline_code(paragraph, line.strip())
        index += 1


def main() -> int:
    """Create the QA flow Word document from a Markdown source path."""
    if len(sys.argv) != 3:
        print("Usage: generate_qa_flow_docx.py <source.md> <output.docx>", file=sys.stderr)
        return 2

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    markdown = source.read_text(encoding="utf-8")

    doc = Document()
    style_document(doc)
    add_title_block(doc)
    doc.add_section(WD_SECTION.CONTINUOUS)
    add_markdown_to_docx(doc, markdown)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Generated {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
